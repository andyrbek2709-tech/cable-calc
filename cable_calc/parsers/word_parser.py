"""Word-парсер кабельных журналов (.docx)."""
import os
from typing import List, Optional
import docx
from .models import CableJournalRow, ParseResult
from .pdf_parser import _row_from_cells
from .utils import clean_ocr


def parse_word(path: str) -> ParseResult:
    result = ParseResult(source_file=os.path.basename(path))
    row_num = 1
    doc = docx.Document(path)

    for table in doc.tables:
        for row in table.rows:
            cells = [clean_ocr(cell.text) for cell in row.cells]
            r = _row_from_cells(cells, row_num)
            if r:
                result.rows.append(r)
                row_num += 1
            else:
                result.skipped_count += 1

    result.total_pages = len(doc.tables)
    result.parsed_count = len(result.rows)
    return result
