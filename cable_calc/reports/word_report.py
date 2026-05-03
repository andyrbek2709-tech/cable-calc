"""Генератор Word-отчёта (docx) с методологией."""
import io
from dataclasses import asdict
from typing import List

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def build_word_report(results: list) -> bytes:
    doc = Document()

    # Поля страницы
    for section in doc.sections:
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)

    # Заголовок
    h = doc.add_heading("Расчёт кабелей до 1кВ", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.runs[0]
    run.font.size = Pt(14)

    sub = doc.add_paragraph("Методика: МЭК 60364-5-52 (IEC 60364-5-52)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    # Сводная таблица
    doc.add_heading("Сводная таблица результатов", 1)
    hdrs = ["№", "Обозначение", "Сечение,мм²", "Iр,А", "Iдоп,А", "ΔU,%", "CB,А", "IКЗ1ф,А", "Статус"]
    tbl = doc.add_table(rows=1, cols=len(hdrs))
    tbl.style = "Table Grid"

    hdr_row = tbl.rows[0]
    for i, h in enumerate(hdrs):
        cell = hdr_row.cells[i]
        cell.text = h
        _set_cell_bg(cell, "1F3864")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    for idx, res in enumerate(results, 1):
        r = res if isinstance(res, dict) else asdict(res)
        status = r.get("status", "?")
        row = tbl.add_row()
        vals = [
            str(idx),
            r.get("line_id", ""),
            str(r.get("section_mm2", "")),
            f"{r.get('i_calc_a', 0):.1f}",
            f"{r.get('i_allowable_a', 0):.1f}",
            f"{r.get('delta_u_pct', 0):.2f}",
            str(r.get("cb_rating_a", "")),
            f"{r.get('i_kz_1ph_a', 0):.0f}",
            status,
        ]
        bg = "E2EFDA" if status == "OK" else ("FCE4D6" if status == "ERROR" else "FFEB9C")
        for ci, v in enumerate(vals):
            cell = row.cells[ci]
            cell.text = v
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # Детальная методология по каждому кабелю
    doc.add_heading("Детальная методология", 1)

    for idx, res in enumerate(results, 1):
        r = res if isinstance(res, dict) else asdict(res)
        doc.add_heading(
            f"{idx}. {r.get('line_id','') or 'Кабель '+str(idx)} — {r.get('status','')}",
            2
        )
        methodology = r.get("methodology", {})
        if not methodology:
            doc.add_paragraph("Методология недоступна.")
            continue

        mtbl = doc.add_table(rows=1, cols=2)
        mtbl.style = "Table Grid"
        mhdr = mtbl.rows[0]
        for ci, txt in enumerate(["Параметр", "Значение"]):
            mhdr.cells[ci].text = txt
            _set_cell_bg(mhdr.cells[ci], "D9D9D9")
            if mhdr.cells[ci].paragraphs[0].runs:
                mhdr.cells[ci].paragraphs[0].runs[0].font.bold = True
                mhdr.cells[ci].paragraphs[0].runs[0].font.size = Pt(9)

        for key, val in methodology.items():
            mrow = mtbl.add_row()
            mrow.cells[0].text = str(key)
            mrow.cells[1].text = str(val)
            for ci in range(2):
                p = mrow.cells[ci].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if p.runs:
                    p.runs[0].font.size = Pt(9)

        # Подсказки
        hints = r.get("hints", [])
        if hints:
            p = doc.add_paragraph()
            run = p.add_run("Рекомендации: ")
            run.font.bold = True
            run.font.size = Pt(9)
            for hint in hints:
                p2 = doc.add_paragraph(f"• {hint}", style="List Bullet")
                p2.runs[0].font.size = Pt(9)
                p2.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
